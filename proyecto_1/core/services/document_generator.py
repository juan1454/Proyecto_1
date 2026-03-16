from docx import Document

from pathlib import Path
from django.conf import settings
from datetime import datetime


def generar_nombre_documento(numero_id):
    """
    Genera un nombre único para el documento
    """

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    return f"word_{numero_id}_{timestamp}.docx"


def generar_word_desde_plantilla(plantilla, datos, numero_id, usuario):
    """
    Genera un documento Word a partir de una plantilla
    """

    # Cargar plantilla
    ruta_plantilla = Path(settings.MEDIA_ROOT) / plantilla.archivo.name
    doc = Document(ruta_plantilla)

    # Reemplazo en párrafos
    for p in doc.paragraphs:
        for key, value in datos.items():
            if key in p.text:
                p.text = p.text.replace(key, str(value))

    # Reemplazo en tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for key, value in datos.items():
                        if key in p.text:
                            p.text = p.text.replace(key, str(value))

    # Carpeta de salida
    salida_dir = Path(settings.MEDIA_ROOT) / "documentos_generados" / usuario
    salida_dir.mkdir(parents=True, exist_ok=True)

    # Nombre automático
    nombre_archivo = generar_nombre_documento(numero_id)
    ruta_salida = salida_dir / nombre_archivo

    doc.save(ruta_salida)

    return ruta_salida


